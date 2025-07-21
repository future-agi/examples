import os
import re
import logging
from typing import Dict, List, Tuple, Optional
from datetime import datetime
import mimetypes

# Document processing imports (will be installed when packages are ready)
try:
    import PyPDF2
    from docx import Document as DocxDocument
    from bs4 import BeautifulSoup
    import requests
    from youtube_transcript_api import YouTubeTranscriptApi
    import whisper
    import speech_recognition as sr
    from pydub import AudioSegment
except ImportError as e:
    logging.warning(f"Some document processing libraries not available: {e}")

class DocumentProcessor:
    """
    Service for processing various document types and extracting text content
    """
    
    def __init__(self):
        self.supported_types = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_txt,
            'md': self._process_markdown,
            'url': self._process_url,
            'youtube': self._process_youtube,
            'audio': self._process_audio
        }
        
        self.mime_type_mapping = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt',
            'text/markdown': 'md',
            'audio/mpeg': 'audio',
            'audio/wav': 'audio',
            'audio/x-wav': 'audio',
            'audio/mp4': 'audio',
            'audio/aac': 'audio',
            'audio/ogg': 'audio',
            'audio/flac': 'audio',
            'video/mp4': 'audio'
        }
    
    def detect_document_type(self, file_path: str = None, url: str = None, mime_type: str = None) -> str:
        """
        Detect the document type based on file extension, URL, or MIME type
        """
        if url:
            if 'youtube.com' in url or 'youtu.be' in url:
                return 'youtube'
            else:
                return 'url'
        
        if mime_type and mime_type in self.mime_type_mapping:
            return self.mime_type_mapping[mime_type]
        
        if file_path:
            _, ext = os.path.splitext(file_path.lower())
            ext_mapping = {
                '.pdf': 'pdf',
                '.docx': 'docx',
                '.doc': 'docx',  # Will try to process as docx
                '.txt': 'txt',
                '.md': 'md',
                '.markdown': 'md',
                '.mp3': 'audio',
                '.wav': 'audio',
                '.m4a': 'audio',
                '.aac': 'audio',
                '.ogg': 'audio',
                '.flac': 'audio',
                '.mp4': 'audio'
            }
            return ext_mapping.get(ext, 'txt')  # Default to txt
        
        return 'txt'
    
    def process_document(self, file_path: str = None, url: str = None, 
                        content: str = None, doc_type: str = None) -> Dict:
        """
        Main method to process a document and extract text content
        
        Returns:
            Dict containing:
            - text: Extracted text content
            - metadata: Document metadata (title, author, pages, etc.)
            - chunks: Text split into manageable chunks
            - processing_stats: Statistics about processing
        """
        start_time = datetime.now()
        
        try:
            # Determine document type
            if not doc_type:
                doc_type = self.detect_document_type(file_path, url)
            
            # Process based on type
            if doc_type in self.supported_types:
                if content:
                    # Direct text content
                    result = self._process_text_content(content)
                else:
                    result = self.supported_types[doc_type](file_path, url)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
            
            # Add processing statistics
            processing_time = (datetime.now() - start_time).total_seconds()
            result['processing_stats'] = {
                'processing_time': processing_time,
                'document_type': doc_type,
                'chunks_created': len(result.get('chunks', [])),
                'word_count': len(result.get('text', '').split()),
                'character_count': len(result.get('text', '')),
                'processed_at': datetime.now().isoformat()
            }
            
            return result
            
        except Exception as e:
            logging.error(f"Error processing document: {str(e)}")
            return {
                'text': '',
                'metadata': {},
                'chunks': [],
                'processing_stats': {
                    'error': str(e),
                    'processing_time': (datetime.now() - start_time).total_seconds(),
                    'processed_at': datetime.now().isoformat()
                }
            }
    
    def _process_pdf(self, file_path: str, url: str = None) -> Dict:
        """Process PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'creator': pdf_reader.metadata.get('/Creator', '') if pdf_reader.metadata else '',
                    'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')) if pdf_reader.metadata else ''
                }
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Page {page_num + 1}]\\n{page_text}")
                    except Exception as e:
                        logging.warning(f"Error extracting text from page {page_num + 1}: {e}")
                
                full_text = "\\n\\n".join(text_content)
                
                return {
                    'text': full_text,
                    'metadata': metadata,
                    'chunks': self._chunk_text(full_text)
                }
                
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def _process_docx(self, file_path: str, url: str = None) -> Dict:
        """Process DOCX files"""
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'paragraphs': len(doc.paragraphs)
            }
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            full_text = "\\n\\n".join(text_content)
            
            return {
                'text': full_text,
                'metadata': metadata,
                'chunks': self._chunk_text(full_text)
            }
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def _process_txt(self, file_path: str, url: str = None) -> Dict:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return self._process_text_content(content)
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    return self._process_text_content(content)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with any supported encoding")
    
    def _process_markdown(self, file_path: str, url: str = None) -> Dict:
        """Process Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Basic markdown processing - remove markdown syntax for plain text
            # Remove headers
            content = re.sub(r'^#{1,6}\\s+', '', content, flags=re.MULTILINE)
            # Remove bold/italic
            content = re.sub(r'\\*\\*([^*]+)\\*\\*', r'\\1', content)
            content = re.sub(r'\\*([^*]+)\\*', r'\\1', content)
            # Remove links but keep text
            content = re.sub(r'\\[([^\\]]+)\\]\\([^)]+\\)', r'\\1', content)
            # Remove code blocks
            content = re.sub(r'```[^`]*```', '', content, flags=re.DOTALL)
            content = re.sub(r'`([^`]+)`', r'\\1', content)
            
            return self._process_text_content(content)
            
        except Exception as e:
            raise Exception(f"Error processing Markdown: {str(e)}")
    
    def _process_url(self, file_path: str = None, url: str = None) -> Dict:
        """Process web URLs"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract metadata
            title = soup.find('title')
            title_text = title.get_text().strip() if title else ''
            
            # Try to find meta description
            description_meta = soup.find('meta', attrs={'name': 'description'})
            description = description_meta.get('content', '') if description_meta else ''
            
            metadata = {
                'title': title_text,
                'description': description,
                'url': url,
                'domain': requests.utils.urlparse(url).netloc
            }
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()
            
            # Extract text content
            text_content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = '\\n'.join(chunk for chunk in chunks if chunk)
            
            return {
                'text': text_content,
                'metadata': metadata,
                'chunks': self._chunk_text(text_content)
            }
            
        except Exception as e:
            raise Exception(f"Error processing URL: {str(e)}")
    
    def _process_youtube(self, file_path: str = None, url: str = None) -> Dict:
        """Process YouTube videos by extracting transcripts"""
        try:
            # Extract video ID from URL
            video_id = self._extract_youtube_id(url)
            if not video_id:
                raise ValueError("Invalid YouTube URL")
            
            # Get transcript
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
            
            # Combine transcript text
            transcript_text = []
            for entry in transcript_list:
                timestamp = self._format_timestamp(entry['start'])
                text = entry['text']
                transcript_text.append(f"[{timestamp}] {text}")
            
            full_text = "\\n".join(transcript_text)
            
            metadata = {
                'video_id': video_id,
                'url': url,
                'duration': transcript_list[-1]['start'] + transcript_list[-1]['duration'] if transcript_list else 0,
                'transcript_entries': len(transcript_list)
            }
            
            return {
                'text': full_text,
                'metadata': metadata,
                'chunks': self._chunk_text(full_text)
            }
            
        except Exception as e:
            raise Exception(f"Error processing YouTube video: {str(e)}")
    
    def _process_text_content(self, content: str) -> Dict:
        """Process direct text content"""
        metadata = {
            'word_count': len(content.split()),
            'character_count': len(content),
            'line_count': len(content.splitlines())
        }
        
        return {
            'text': content,
            'metadata': metadata,
            'chunks': self._chunk_text(content)
        }
    
    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict]:
        """
        Split text into overlapping chunks for better processing
        """
        if not text:
            return []
        
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunks.append({
                'text': chunk_text,
                'start_word': i,
                'end_word': min(i + chunk_size, len(words)),
                'word_count': len(chunk_words),
                'character_count': len(chunk_text)
            })
            
            # Break if we've reached the end
            if i + chunk_size >= len(words):
                break
        
        return chunks
    
    def _extract_youtube_id(self, url: str) -> Optional[str]:
        """Extract YouTube video ID from URL"""
        patterns = [
            r'(?:youtube\\.com/watch\\?v=|youtu\\.be/)([^&\\n?#]+)',
            r'youtube\\.com/embed/([^&\\n?#]+)',
            r'youtube\\.com/v/([^&\\n?#]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        
        return None
    
    def _format_timestamp(self, seconds: float) -> str:
        """Format seconds to MM:SS or HH:MM:SS format"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        seconds = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        else:
            return f"{minutes:02d}:{seconds:02d}"
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported document types"""
        return list(self.supported_types.keys())
    
    def validate_file_size(self, file_path: str, max_size_mb: int = 200) -> Tuple[bool, str]:
        """Validate file size"""
        try:
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            
            if file_size > max_size_bytes:
                return False, f"File size ({file_size / (1024*1024):.1f}MB) exceeds maximum allowed size ({max_size_mb}MB)"
            
            return True, "OK"
            
        except Exception as e:
            return False, f"Error checking file size: {str(e)}"
    
    def _process_audio(self, file_path: str, url: str = None) -> Dict:
        """Process audio files using Whisper for transcription with optimizations"""
        try:
            # Check if required audio libraries are available
            try:
                import whisper
                import speech_recognition as sr
                from pydub import AudioSegment
                from src.services.audio_optimizer import AudioOptimizer
            except ImportError as import_err:
                raise Exception(f"Audio processing libraries not available: {import_err}")
            
            # Check if ffmpeg is available (required by pydub)
            try:
                # Test pydub functionality first
                test_audio = AudioSegment.silent(duration=1000)  # 1 second of silence
            except Exception as pydub_err:
                raise Exception(f"FFmpeg not available or pydub error: {pydub_err}")
            
            # Get file size for optimization
            file_size = os.path.getsize(file_path)
            whisper_settings = AudioOptimizer.optimize_whisper_settings(file_size)
            
            # Load the Whisper model based on file size
            logging.info(f"Loading Whisper model '{whisper_settings['model']}' for audio file: {file_path}")
            cache_dir = os.environ.get('HF_HOME', '/app/model_cache')
            model = whisper.load_model(whisper_settings['model'], download_root=cache_dir)
            
            # Get audio file info
            logging.info(f"Reading audio file: {file_path}")
            audio = AudioSegment.from_file(file_path)
            duration_seconds = len(audio) / 1000.0
            
            logging.info(f"Audio duration: {duration_seconds:.2f} seconds, Size: {file_size/(1024*1024):.2f} MB")
            
            # Convert to WAV if needed (Whisper works best with WAV)
            temp_wav_path = None
            if not file_path.lower().endswith('.wav'):
                temp_wav_path = file_path.rsplit('.', 1)[0] + '_temp.wav'
                logging.info(f"Converting audio to WAV: {temp_wav_path}")
                # Use 16kHz sampling rate for better performance
                audio = audio.set_frame_rate(16000)
                audio.export(temp_wav_path, format="wav")
                transcription_path = temp_wav_path
            else:
                transcription_path = file_path
            
            # Transcribe the audio with optimized settings
            logging.info(f"Starting transcription of: {transcription_path}")
            result = model.transcribe(
                transcription_path,
                temperature=whisper_settings['temperature'],
                compression_ratio_threshold=whisper_settings['compression_ratio_threshold'],
                logprob_threshold=whisper_settings['logprob_threshold'],
                no_speech_threshold=whisper_settings['no_speech_threshold']
            )
            logging.info(f"Transcription completed. Detected language: {result.get('language', 'unknown')}")
            
            # Clean up temporary file
            if temp_wav_path and os.path.exists(temp_wav_path):
                os.remove(temp_wav_path)
                logging.info(f"Cleaned up temporary file: {temp_wav_path}")
            
            # Process transcript for better Q&A
            clean_text, segments_with_timestamps = AudioOptimizer.process_whisper_transcript(result)
            
            # Create optimized chunks
            chunks = AudioOptimizer.create_smart_chunks(
                clean_text, 
                segments_with_timestamps,
                chunk_size=500,  # Smaller chunks for audio
                overlap=50
            )
            
            # Create comprehensive metadata
            metadata = AudioOptimizer.create_audio_metadata(
                file_path, result, clean_text, segments_with_timestamps
            )
            
            # Add additional metadata
            metadata.update({
                'sample_rate': audio.frame_rate,
                'channels': audio.channels,
                'whisper_model': whisper_settings['model']
            })
            
            logging.info(f"Audio processing completed. Clean transcript: {len(clean_text)} chars, {len(chunks)} chunks")
            
            # Store both clean text and timestamped version
            return {
                'text': clean_text,  # Clean text for vector search
                'text_with_timestamps': "\n".join([
                    f"[{seg['start_formatted']} - {seg['end_formatted']}] {seg['text']}"
                    for seg in segments_with_timestamps
                ]),  # Timestamped version for reference
                'metadata': metadata,
                'chunks': chunks,
                'segments': segments_with_timestamps  # Store segments for future use
            }
            
        except Exception as e:
            logging.error(f"Error processing audio file {file_path}: {str(e)}")
            raise Exception(f"Error processing audio file: {str(e)}")

